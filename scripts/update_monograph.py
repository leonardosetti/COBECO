from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "Monografia e Demais Docs"
MAIN = DOC_DIR / "COBECO_Monografia.docx"
LEGACY = DOC_DIR / "COBECO_objetivos_escopo_requisitos_ABNT.docx"

functional = [
    ("RF01", "Permitir cadastro com nome, e-mail único, senha e consentimento."),
    ("RF02", "Autenticar por e-mail e senha e direcionar à área logada."),
    ("RF03", "Encerrar a sessão com segurança."),
    ("RF04", "Solicitar e concluir a redefinição de senha."),
    ("RF05", "Proteger todas as rotas da plataforma."),
    ("RF06", "Criar listas nomeadas vinculadas a uma categoria."),
    ("RF07", "Incluir item com produto normalizado ou descrição livre, categoria e quantidade."),
    ("RF08", "Editar o nome da lista e editar ou remover seus itens."),
    ("RF09", "Duplicar uma lista e seus itens."),
    ("RF10", "Excluir uma lista após confirmação."),
    ("RF11", "Exibir as listas do usuário, seus itens e datas."),
    ("RF12", "Solicitar a cotação de todos os itens de uma lista."),
    ("RF13", "Agrupar fornecedores com exatamente o mesmo conjunto de itens disponíveis e exibir o percentual de cobertura."),
    ("RF14", "Calcular, por grupo, totais individuais e o menor custo consolidado por item."),
    ("RF15", "Destacar visualmente o melhor grupo calculado."),
    ("RF16", "Listar os produtos ausentes em cada grupo."),
    ("RF17", "Permitir incluir ou excluir fornecedores da comparação."),
    ("RF18", "Apresentar painel com fornecedores, cobertura, preços, ausências e orçamento consolidado."),
    ("RF19", "Persistir cotações, data, grupos e resultado destacado."),
    ("RF20", "Visualizar e reabrir cotações anteriores."),
    ("RF21", "Excluir registros do histórico pertencentes ao usuário."),
    ("RF22", "Informar quando nenhum fornecedor ou oferta válida for encontrado."),
    ("RF23", "Sinalizar produto não mapeado ou indisponível."),
    ("RF24", "Validar nome, quantidade, categoria e seleção antes da cotação."),
]

non_functional = [
    ("RNF01", "Separar domínio puro, serviços, adapters de persistência e framework web."),
    ("RNF02", "Manter API REST regida por contrato OpenAPI e adotar API-First nas novas rotas."),
    ("RNF03", "Usar PostgreSQL com transações ACID e migrations versionadas."),
    ("RNF04", "Disponibilizar frontend, API e PostgreSQL como três serviços Docker."),
    ("RNF05", "Descrever regras críticas em Gherkin e executá-las com Cucumber."),
    ("RNF06", "Cobrir domínio e serviços com testes Vitest."),
    ("RNF07", "Cobrir cadastro, login, lista e cotação com Playwright."),
    ("RNF08", "Executar lint, testes e build no GitHub Actions."),
    ("RNF09", "Otimizar a interface para desktop a partir de 1024 × 768."),
    ("RNF10", "Exibir estados de carregamento e erros compreensíveis."),
    ("RNF11", "Garantir navegação por teclado, foco, contraste, semântica e rótulos."),
    ("RNF12", "Destacar o melhor orçamento sem depender somente de cor."),
    ("RNF13", "Armazenar senhas somente com hash Argon2."),
    ("RNF14", "Usar HTTPS no ambiente de produção."),
    ("RNF15", "Expirar a sessão após 30 minutos sem renovação ou atividade."),
    ("RNF16", "Indexar chaves usadas no catálogo, histórico e relacionamentos."),
    ("RNF17", "Servir a UI estática independentemente da disponibilidade da API."),
    ("RNF18", "Utilizar tecnologias FOSS ou de uso gratuito."),
    ("RNF19", "Implementar o backend em Node.js, TypeScript e Express."),
    ("RNF20", "Implementar o frontend em React 18 e TypeScript, sem jQuery."),
    ("RNF21", "Usar PostgreSQL 16 e Prisma ORM 5."),
    ("RNF22", "Versionar no GitHub e adotar GitHub Flow."),
]

desirable = [
    ("RD01", "Autenticação por contas de terceiros."),
    ("RD02", "Alerta de preço abaixo de limite definido."),
    ("RD03", "Exportação avançada dos resultados em PDF ou planilha."),
    ("RD04", "Gráficos de variação histórica de preços."),
    ("RD05", "Ampliação do compartilhamento e colaboração em listas."),
    ("RD06", "Novas categorias e painel administrativo do catálogo."),
]


def set_table(table, rows):
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    table.cell(0, 0).text = "Código"
    table.cell(0, 1).text = "Descrição do requisito"
    for index, (code, description) in enumerate(rows, 1):
        table.cell(index, 0).text = code
        table.cell(index, 1).text = description


doc = Document(MAIN)
paragraphs = doc.paragraphs

# Em execuções posteriores, atualiza apenas o trecho sujeito a mudanças de versão.
if paragraphs[89].text.startswith("O COBECO (Cotação de Bens de Consumo) é uma aplicação web destinada"):
    paragraphs[151].text = (
        "Foram adotados Node.js 20, Express 4.22, TypeScript 5.9, Zod 3.25, JWT, Argon2 0.44, Prisma 5.22 e PostgreSQL 16 no backend; "
        "React 18.3, React Router 7.18, Vite 8.2 e Tailwind CSS 3.4 no frontend; Vitest 4.1, Cucumber 13.2 e Playwright 1.62 para "
        "qualidade; OpenAPI 3.0.3 para o contrato; e Docker Compose com Nginx para os três serviços. Todas as tecnologias são FOSS ou de uso gratuito."
    )
    doc.save(MAIN)
    print(f"Versões atualizadas: {MAIN}")
    raise SystemExit(0)

paragraphs[89].text = (
    "O COBECO (Cotação de Bens de Consumo) é uma aplicação web destinada a usuários autenticados que desejam organizar listas "
    "e comparar fornecedores de forma transparente. O núcleo do MVP é o motor de paridade: para cada lista, o sistema identifica "
    "os itens disponíveis em cada fornecedor, agrupa fornecedores com o mesmo perfil de cobertura, calcula os totais e evidencia "
    "produtos ausentes e o orçamento consolidado. Os preços do MVP provêm de catálogo controlado e persistido no PostgreSQL, o "
    "que assegura uma demonstração determinística e independente de serviços externos. Listas e cotações permanecem disponíveis "
    "no histórico do usuário, com autenticação, validação e tratamento de exceções."
)
paragraphs[93].text = (
    "O escopo inclui cadastro, login, logout e recuperação de senha; gestão completa de listas categorizadas e seus itens; catálogo "
    "normalizado de categorias, produtos, fornecedores e ofertas; seleção dos fornecedores; cotação de lista; agrupamento por perfil "
    "idêntico de cobertura; cálculo do percentual, dos itens ausentes, dos totais por fornecedor e do menor custo consolidado; painel "
    "comparativo; persistência e reabertura do histórico; exportação CSV e impressão; contrato OpenAPI; e execução em contêineres "
    "separados para interface, API e banco. Não integram o MVP compras, pagamentos, garantia de preço em tempo real, aplicativo móvel "
    "nativo, autenticação social ou dependência obrigatória de APIs externas. Integrações já desenvolvidas são preservadas, mas ficam "
    "desativadas por padrão como evolução pós-MVP."
)

for table, rows in zip(doc.tables[:5], [functional[:5], functional[5:11], functional[11:18], functional[18:21], functional[21:24]]):
    set_table(table, rows)

for table, rows in zip(doc.tables[5:9], [non_functional[:4], non_functional[4:8], non_functional[8:12], non_functional[12:22]]):
    set_table(table, rows)

set_table(doc.tables[9], desirable)

for index, value in {
    102: "3.1.1 Autenticação e controle de acesso",
    106: "3.1.2 Gestão de listas e produtos",
    109: "3.1.3 Cotação por paridade",
    113: "3.1.4 Histórico e persistência",
    116: "3.1.5 Validação e exceções",
    124: "3.2.1 Arquitetura e infraestrutura",
    127: "3.2.2 Testes e qualidade",
    130: "3.2.3 Experiência e acessibilidade",
    134: "3.2.4 Segurança, desempenho e stack",
    145: "4 TECNOLOGIAS, MODELAGEM E PROTÓTIPOS",
    149: "4.1 Tecnologias adotadas",
    153: "4.2 Modelagem, protótipos e qualidade",
    157: "5 CONSIDERAÇÕES FINAIS",
}.items():
    paragraphs[index].text = value

paragraphs[147].text = (
    "A solução foi implementada como monorepo TypeScript. O domínio de disponibilidade, agrupamento e orçamento é independente de "
    "Express e Prisma. A API expõe contratos REST documentados em OpenAPI e persiste dados com Prisma/PostgreSQL. A interface React "
    "permite selecionar fornecedores e apresenta grupos, cobertura, ausências, totais individuais e destaque textual do resultado."
)
paragraphs[151].text = (
    "Foram adotados Node.js 20, Express 4.22, TypeScript 5.9, Zod 3.25, JWT, Argon2 0.44, Prisma 5.22 e PostgreSQL 16 no backend; React 18.3, "
    "React Router 7.18, Vite 8.2 e Tailwind CSS 3.4 no frontend; Vitest 4.1, Cucumber 13.2 e Playwright 1.62 para qualidade; OpenAPI 3.0.3 para o contrato; e Docker "
    "Compose com Nginx para os três serviços. Todas as tecnologias são FOSS ou de uso gratuito."
)
paragraphs[155].text = (
    "O DER deriva do schema Prisma e inclui usuários, listas, itens, categorias, produtos, fornecedores, ofertas e cotações. O diagrama "
    "de casos de uso foi revisado para representar o catálogo persistido e o agrupamento por paridade. Os protótipos de baixa "
    "fidelidade estão nos arquivos Excalidraw e o protótipo de alta fidelidade é a interface React executável. O cenário A–H gera "
    "quatro grupos e é validado por teste unitário, integração HTTP e cenário Gherkin; o fluxo crítico foi validado no Chromium."
)
paragraphs[159].text = (
    "O projeto alcança o objetivo de transformar uma lista de bens de consumo em uma comparação clara de cobertura e custo. Ao "
    "agrupar fornecedores pelo conjunto real de itens disponíveis, o COBECO evita que percentuais iguais ocultem perfis distintos e "
    "torna explícitas as pendências de cada alternativa. O catálogo seedado reduz riscos externos, enquanto PostgreSQL, OpenAPI, "
    "contêineres e testes tornam a solução reproduzível. Como continuidade, recomenda-se ampliar o catálogo, criar sua administração, "
    "automatizar auditorias de acessibilidade e publicar a aplicação atrás de HTTPS."
)
paragraphs[160].text = ""

# Atualiza o sumário textual; a paginação deve ser recalculada pelo Word na exportação final.
toc = {
    62: "1 OBJETIVO DO PROJETO", 63: "2 ESCOPO", 64: "3 REQUISITOS", 65: "3.1 Funcionais",
    66: "3.1.1 Autenticação e controle de acesso", 67: "3.1.2 Gestão de listas e produtos",
    68: "3.1.3 Cotação por paridade", 69: "3.1.4 Histórico e persistência",
    70: "3.1.5 Validação e exceções", 71: "3.2 Não-Funcionais",
    72: "3.2.1 Arquitetura e infraestrutura", 73: "3.2.2 Testes e qualidade",
    74: "3.2.3 Experiência e acessibilidade", 75: "3.2.4 Segurança, desempenho e stack",
    76: "3.3 Desejáveis", 77: "4 TECNOLOGIAS, MODELAGEM E PROTÓTIPOS",
    78: "4.1 Tecnologias adotadas", 79: "4.2 Modelagem, protótipos e qualidade",
    80: "5 CONSIDERAÇÕES FINAIS", 81: "REFERÊNCIAS",
}
for index, value in toc.items():
    paragraphs[index].text = value
for index in range(82, 86):
    paragraphs[index].text = ""

# Substitui o conteúdo-modelo de referências e remove apêndices/anexos vazios.
paragraphs[164].text = "OPENAPI INITIATIVE. OpenAPI Specification 3.0.3. 2020."
paragraphs[166].text = "POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL 16 Documentation. 2026."
paragraphs[167].text = "PRISMA DATA. Prisma ORM Documentation. 2026."
paragraphs[169].text = "META PLATFORMS. React Documentation. 2026."
for index in list(range(170, 188)) + list(range(213, 270)):
    paragraphs[index]._element.getparent().remove(paragraphs[index]._element)

doc.core_properties.title = "COBECO — Cotação de Bens de Consumo"
doc.core_properties.subject = "Documentação do MVP e motor de paridade"
doc.save(MAIN)

# O documento menor continua disponível, mas deixa explícito que foi superado.
legacy = Document(LEGACY)
if not legacy.paragraphs[0].text.startswith("DOCUMENTO HISTÓRICO"):
    warning = legacy.paragraphs[0].insert_paragraph_before(
        "DOCUMENTO HISTÓRICO — conteúdo superado em 30/08/2026. Utilize COBECO_Monografia.docx e prj_docs/md/objetivos.md."
    )
    warning.style = legacy.paragraphs[0].style
legacy.save(LEGACY)

print(f"Atualizado: {MAIN}")
print(f"Marcado como histórico: {LEGACY}")
